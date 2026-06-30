"""
collections.py  —  Capstone Project 8 · RAG Ingestion Pipeline
=========================================================================

PURPOSE
-------
This script converts static knowledge files (TXT / PDF / DOCX) into dense
vector embeddings and stores them in ChromaDB so that LangGraph agents can
retrieve semantically relevant context at runtime.

The pipeline works in four stages:
  1. LOAD   — Read raw text from disk (.txt / .pdf / .docx)
  2. CHUNK  — Split text into overlapping windows (avoids truncated context)
  3. EMBED  — Encode each chunk into a 384-dim float32 vector (MiniLM model)
  4. UPSERT — Write (text, vector, metadata) triples into ChromaDB

EMBEDDING MODEL
---------------
Model:   all-MiniLM-L6-v2  (from sentence-transformers)
Dims:    384 float32 values per chunk
Speed:   ~14 000 sentences/sec on CPU (no GPU required)
Source:  https://huggingface.co/sentence-transformers/all-MiniLM-L6-v2

ChromaDB calls SentenceTransformerEmbeddingFunction, which wraps the model.
When you call collection.upsert(documents=[...]), ChromaDB internally calls
  vectors = embedding_fn(documents)
and stores both the raw text and the resulting float32 vectors.
At query time, the query string is embedded the same way, and ChromaDB
computes cosine distance between the query vector and every stored vector,
returning the top-N closest chunks.

DISTANCE METRIC
---------------
All three collections use "cosine" space (hnsw:space = "cosine") with
normalize_embeddings=True — so "lower distance = more similar" and the
distance range is [0.0, 2.0] (0.0 = identical, 2.0 = opposite).

THREE COLLECTIONS (vs one monolithic collection in src/rag/utils.py)
-----------------------------------------------------------------
This script creates NAMED, domain-specific collections, each containing
different types of knowledge:

  historical_precedents   — Case studies of past supply-chain crises
                            (Chunk size: 700 chars — short crisp events)
  export_control_corpus   — Policy docs (BIS rules, CHIPS Act, sanctions)
                            (Chunk size: 900 chars — legal text needs context)
  india_sourcing_corpus   — India sourcing / PLI scheme reports
                            (Chunk size: 600 chars — structured report bullets)

These are SEPARATE from the single 'electronics_supply_chain_knowledge'
collection built by src/rag/utils.py (which ingests the Excel workbook).
Agents query by name:
  - Agent 4 (Risk Classifier)    → historical_precedents + export_control_corpus
  - Agent 7 (Mitigation)         → all three collections
  - Agent 2 (News Event)         → historical_precedents only
  - Agent 3 (Weather Monitoring) → historical_precedents only (severity >= 7)

DIRECTORY LAYOUT EXPECTED
-------------------------
  data/raw/RAG_data/
    historical_precedents/     ← .txt / .pdf / .docx files go here
    export_control_corpus/
    india_sourcing_corpus/

OUTPUT
------
  outputs/chromadb/            ← ChromaDB SQLite + vector index files

USAGE
-----
  python scripts/build_rag_collections.py               # incremental upsert
  python scripts/build_rag_collections.py --flush       # wipe + full rebuild
  python scripts/build_rag_collections.py --query "chip shortage Taiwan"
  python scripts/build_rag_collections.py --collection historical_precedents
"""

from __future__ import annotations

import argparse
import hashlib
import re
import shutil
import sys
from pathlib import Path
from typing import Any, Dict, List, Tuple

# ---------------------------------------------------------------------------
# Global configuration — change these if you swap the model or tune chunking
# ---------------------------------------------------------------------------

# Project root is two levels up from this file:
#   __file__        → src/rag/collections.py
#   .parent         → src/rag/
#   .parent.parent  → src/
#   .parents[2]     → project root  (where data/, outputs/, src/ live)
PROJECT_ROOT = Path(__file__).resolve().parents[2]

# Where raw knowledge files live (one subdirectory per collection)
RAG_DATA_ROOT = PROJECT_ROOT / "data" / "raw" / "RAG_data"

# Where ChromaDB persists its SQLite + HNSW index files (shared with src.rag.utils)
from src.rag.utils import (
    CHROMA_DIR,
    EMBEDDING_MODEL,
    FINETUNED_EMBEDDING_MODEL,
    get_chroma_client,
    get_embedding_model,
    reset_chroma_client,
    resolve_embedding_model_name,
)

# Default chunk parameters (overridden per collection in COLLECTION_CHUNK_CONFIG)
CHUNK_SIZE = 900      # Maximum characters per chunk
CHUNK_OVERLAP = 150   # Characters of overlap between consecutive chunks
                      # — ensures sentences split across a boundary appear
                      #   in at least one complete chunk

# Canonical collection names — agents import and use these strings directly
COLLECTION_NAMES = {
    "historical_precedents": "historical_precedents",
    "export_control_corpus": "export_control_corpus",
    "india_sourcing_corpus": "india_sourcing_corpus",
}

# Per-collection chunk tuning:
# - Larger chunks preserve more context for dense legal / policy text
# - Smaller chunks give sharper retrieval for short event descriptions
COLLECTION_CHUNK_CONFIG: Dict[str, Dict[str, int]] = {
    "historical_precedents": {"size": 700,  "overlap": 100},
    "export_control_corpus": {"size": 900,  "overlap": 150},
    "india_sourcing_corpus": {"size": 600,  "overlap": 100},
}


# ---------------------------------------------------------------------------
# Lazy imports
# — We defer heavy imports so the script gives a clear error if a package
#   is missing, rather than crashing with a stack trace.
# ---------------------------------------------------------------------------

def _import_chromadb():
    """
    Import chromadb and the SentenceTransformer embedding function wrapper.
    ChromaDB uses this wrapper to call the sentence-transformers library,
    which downloads and caches the model on first use (~90 MB).
    """
    try:
        import chromadb
        from chromadb.utils.embedding_functions import SentenceTransformerEmbeddingFunction
        return chromadb, SentenceTransformerEmbeddingFunction
    except ImportError:
        sys.exit("ERROR: chromadb not installed. Run: pip install chromadb")


def _import_pdf():
    """
    Import pypdf for PDF text extraction.
    Returns None if not installed (PDF ingestion will be skipped with a warning).
    """
    try:
        from pypdf import PdfReader
        return PdfReader
    except ImportError:
        return None


def _import_docx():
    """
    Import python-docx for Word document extraction.
    Returns None if not installed (DOCX ingestion will be skipped with a warning).
    """
    try:
        from docx import Document
        return Document
    except ImportError:
        return None


# ---------------------------------------------------------------------------
# Stage 2 — CHUNKING
#
# Why do we chunk at all?
#   LLM embeddings have a token limit (MiniLM caps at 256 tokens / ~1 000 chars).
#   Long documents must be split. Overlap prevents critical context from being
#   silently dropped when a sentence straddles a chunk boundary.
#
# Sliding window algorithm:
#   start = 0
#   while more text:
#     end = start + CHUNK_SIZE
#     try to break at a paragraph / sentence / newline boundary near `end`
#     emit chunk text[start:end]
#     start = end - CHUNK_OVERLAP        ← step back so next chunk overlaps
# ---------------------------------------------------------------------------

def _chunk_text(text: str, size: int = CHUNK_SIZE, overlap: int = CHUNK_OVERLAP) -> List[str]:
    """
    Split `text` into overlapping character windows, breaking at natural
    paragraph / sentence boundaries wherever possible.

    Boundary preference order (most to least preferred):
      1. Double newline  (paragraph break)
      2. Period + newline (end of sentence followed by new line)
      3. Period + space  (end of inline sentence)
      4. Single newline  (line break)

    If no boundary is found in the second half of the window, we split
    at the hard character limit — no chunk is ever longer than `size`.

    Args:
        text:    The full document text.
        size:    Max characters per chunk.
        overlap: Characters to repeat at the start of the next chunk.

    Returns:
        A list of non-empty string chunks.
    """
    # Normalise Windows line endings to Unix newlines
    text = re.sub(r"\r\n?", "\n", text).strip()

    if not text:
        return []

    # Short documents fit in one chunk — no splitting needed
    if len(text) <= size:
        return [text]

    chunks: List[str] = []
    start = 0

    while start < len(text):
        end = min(len(text), start + size)

        # Only try smart boundary detection when we're not yet at the end
        if end < len(text):
            # Search for a natural break in the SECOND HALF of the window
            # (first half is off-limits so chunks can't become tiny)
            search_start = start + size // 2
            for sep in ["\n\n", ".\n", ". ", "\n"]:
                idx = text.rfind(sep, search_start, end)
                if idx > start:
                    end = idx + len(sep)   # include the separator in the chunk
                    break

        chunk = text[start:end].strip()
        if chunk:
            chunks.append(chunk)

        if end >= len(text):
            break   # reached the document end

        # Step back by `overlap` characters so the next chunk shares context
        start = max(start + 1, end - overlap)

    return chunks


# ---------------------------------------------------------------------------
# Stage 1 — LOAD
# Each loader returns plain UTF-8 text. Tables in DOCX are serialised as
# pipe-delimited rows so the semantic content survives the text conversion.
# ---------------------------------------------------------------------------

def _load_txt(path: Path) -> str:
    """Read a plain-text file. 'errors=replace' avoids crashes on bad encodings."""
    return path.read_text(encoding="utf-8", errors="replace")


def _load_pdf(path: Path) -> str:
    """
    Extract text from each page of a PDF using pypdf.
    Pages are joined with double newlines so paragraph context is preserved.
    Empty pages (e.g. cover images) are silently skipped.
    """
    PdfReader = _import_pdf()
    if PdfReader is None:
        print(f"  [SKIP] {path.name} — pypdf not installed")
        return ""
    reader = PdfReader(path)
    pages = []
    for page in reader.pages:
        text = page.extract_text() or ""
        if text.strip():
            pages.append(text)
    return "\n\n".join(pages)


def _load_docx(path: Path) -> str:
    """
    Extract text from a Word document using python-docx.
    - Paragraphs are joined with double newlines.
    - Tables are serialised as "cell1 | cell2 | ..." rows and tagged
      with a [Table N] header so the model can recognise tabular content.
    """
    Document = _import_docx()
    if Document is None:
        print(f"  [SKIP] {path.name} — python-docx not installed")
        return ""
    doc = Document(path)
    parts: List[str] = []

    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text.strip())

    for i, table in enumerate(doc.tables, 1):
        rows = []
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                rows.append(" | ".join(cells))
        if rows:
            parts.append(f"[Table {i}]\n" + "\n".join(rows))

    return "\n\n".join(parts)


def _load_file(path: Path) -> str:
    """Dispatch to the correct loader based on file extension."""
    ext = path.suffix.lower()
    if ext == ".txt":
        return _load_txt(path)
    elif ext == ".pdf":
        return _load_pdf(path)
    elif ext == ".docx":
        return _load_docx(path)
    else:
        return ""   # unsupported — caller will skip with a warning


# ---------------------------------------------------------------------------
# Metadata helpers
# ---------------------------------------------------------------------------

def _doc_id(collection: str, filename: str, chunk_index: int) -> str:
    """
    Generate a stable, deterministic SHA-256 chunk ID.

    Using a content-addressable ID means re-running the builder on the same
    files produces the same IDs → ChromaDB's upsert is idempotent (it
    overwrites the existing vector rather than creating a duplicate).

    We truncate to 32 hex characters (128 bits) — collision probability is
    negligible for a corpus of thousands of chunks.
    """
    key = f"{collection}|{filename}|{chunk_index}".encode()
    return hashlib.sha256(key).hexdigest()[:32]


def _infer_metadata(filename: str, collection: str) -> Dict[str, Any]:
    """
    Derive rich metadata from the filename without reading the file content.
    This metadata is stored alongside each chunk in ChromaDB and can be used
    as a WHERE filter when agents query the collection.

    Fields added:
      collection          — which named collection this chunk belongs to
      source_file         — original filename (for citations in the dashboard)
      domain              — always "electronics_semiconductors"
      dataset_owner       — project attribution
      severity_hint       — inferred risk severity from filename keywords
                            (CRITICAL / HIGH / MEDIUM / LOW)
      disruption_category — supply-chain disruption type from filename keywords
                            (weather / geopolitical / logistics / factory / ...)
      event_year          — year extracted from filename if present (e.g. "2021")
    """
    name = filename.lower()
    meta: Dict[str, Any] = {
        "collection":    collection,
        "source_file":   filename,
        "domain":        "electronics_semiconductors",
        "dataset_owner": "TeamDisruptors_Capstone_Project8",
    }

    # --- Severity inference: scan for high-signal keywords in filename ---
    # Order matters: more specific (higher severity) keywords are checked first.
    severity_keyword_map = [
        ("CRITICAL", ["covid", "chip_shortage", "bis_export", "chips_act"]),
        ("HIGH",     ["red_sea", "foxconn", "gallium", "ukraine", "taiwan_earth",
                      "hbm", "japan_semi", "china_critical", "gao", "crs_"]),
        ("MEDIUM",   ["ever_given", "thailand", "texas_winter", "micron_india",
                      "tata_", "pli_scheme"]),
    ]
    for severity_label, keywords in severity_keyword_map:
        if any(kw in name for kw in keywords):
            meta["severity_hint"] = severity_label
            break
    else:
        # No strong keyword → default to LOW
        meta["severity_hint"] = "LOW"

    # --- Disruption category: what kind of supply-chain event is this? ---
    category_keyword_map = {
        "weather":      ["weather", "storm", "flood", "earthquake", "hurricane"],
        "geopolitical": ["export_control", "bis_", "japan_semi", "china_critical",
                         "gallium", "ukraine", "chips_act", "crs_", "gao_"],
        "logistics":    ["red_sea", "ever_given", "suez"],
        "factory":      ["foxconn", "lockdown"],
        "demand_shock": ["covid", "hbm", "chip_shortage"],
        "india_policy": ["india_semi", "micron_india", "tata_", "pli_",
                         "kaynes", "cg_power", "invest_india"],
    }
    for category, kws in category_keyword_map.items():
        if any(kw in name for kw in kws):
            meta["disruption_category"] = category
            break
    else:
        meta["disruption_category"] = "general"

    # --- Year inference: extract a 4-digit year from the filename if present ---
    for year in range(2010, 2026):
        if str(year) in name:
            meta["event_year"] = year
            break

    return meta


# ---------------------------------------------------------------------------
# Stage 3 + 4 — EMBED + UPSERT (core ingestion loop)
#
# Data flow per file:
#   raw text
#     → _chunk_text()         →  List[str]  (N chunks)
#     → _infer_metadata()     →  Dict        (shared base metadata)
#     → _doc_id() per chunk   →  str ID      (stable hash)
#     → collection.upsert()   →  ChromaDB calls embedding_fn(chunks)
#                                           → 384-dim float32 vectors stored
# ---------------------------------------------------------------------------

def build_collection(
    collection_name: str,
    source_dir: Path,
    chroma_client,
    embedding_fn,
    chunk_cfg: Dict[str, int],
) -> Dict[str, Any]:
    """
    Ingest all .txt / .pdf / .docx files from `source_dir` into a named
    ChromaDB collection, returning a summary dict.

    ChromaDB's upsert() is idempotent: running this function twice on the
    same files is safe — existing chunks are overwritten, not duplicated,
    because we use deterministic SHA-256 IDs.

    Internally, ChromaDB stores each chunk as a triple:
        (id: str,  document: str,  embedding: List[float])
    Metadata is stored separately in a SQLite table within the ChromaDB dir.

    Args:
        collection_name: One of the names defined in COLLECTION_NAMES.
        source_dir:      Directory containing the knowledge files.
        chroma_client:   An open chromadb.PersistentClient instance.
        embedding_fn:    SentenceTransformerEmbeddingFunction wrapping MiniLM.
        chunk_cfg:       Dict with "size" and "overlap" keys for this collection.

    Returns:
        {"collection": str, "chunks": int, "files": int, "source_dir": str}
    """
    print(f"\n{'='*60}")
    print(f"  Building: {collection_name}")
    print(f"  Source:   {source_dir}")
    print(f"{'='*60}")

    if not source_dir.exists():
        print(f"  [WARN] Source directory not found: {source_dir}")
        return {"collection": collection_name, "chunks": 0, "files": 0}

    # Get the collection if it exists, or create it fresh.
    # The embedding_function is stored in the collection; every future
    # query on this collection will use the SAME model automatically.
    # hnsw:space = "cosine" means distances are in [0.0, 2.0]:
    #   0.0 → perfectly identical vectors
    #   1.0 → orthogonal (unrelated)
    #   2.0 → opposite (contradictory)
    collection = chroma_client.get_or_create_collection(
        name=collection_name,
        embedding_function=embedding_fn,
        metadata={
            "hnsw:space":  "cosine",           # similarity metric
            "domain":      "electronics_semiconductors",
            "capstone":    "Project8_Team11",
        },
    )

    # Discover all supported files in the source directory
    supported_extensions = [".txt", ".pdf", ".docx"]
    files = sorted(
        f for f in source_dir.iterdir()
        if f.suffix.lower() in supported_extensions
    )

    # Accumulators — we batch all chunks before upserting to ChromaDB
    # so we can send them in efficient batches of 64.
    all_docs:  List[str]           = []   # raw chunk text
    all_metas: List[Dict[str, Any]] = []  # metadata dict per chunk
    all_ids:   List[str]           = []   # deterministic chunk IDs

    size    = chunk_cfg["size"]
    overlap = chunk_cfg["overlap"]

    # ── Per-file ingestion ────────────────────────────────────────────────────
    for fpath in files:
        print(f"  Reading: {fpath.name}")

        # Stage 1: Load raw text from disk
        text = _load_file(fpath)
        if not text.strip():
            print(f"    [WARN] Empty content, skipping.")
            continue

        # Stage 2: Split into overlapping character windows
        chunks = _chunk_text(text, size=size, overlap=overlap)

        # Build the base metadata dict from the filename (no file reading needed)
        base_meta = _infer_metadata(fpath.name, collection_name)
        base_meta["total_chunks"] = len(chunks)  # handy for debugging

        # Stage 3 prep: assign a stable ID and per-chunk metadata to each chunk.
        # The actual embedding (Stage 3) happens inside collection.upsert() below.
        for idx, chunk in enumerate(chunks):
            doc_id     = _doc_id(collection_name, fpath.name, idx)
            chunk_meta = {**base_meta, "chunk_index": idx}

            all_docs.append(chunk)
            all_metas.append(chunk_meta)
            all_ids.append(doc_id)

        print(f"    → {len(chunks)} chunks  (size={size}, overlap={overlap})")

    if not all_docs:
        print(f"  [WARN] No documents ingested for {collection_name}")
        return {"collection": collection_name, "chunks": 0, "files": len(files)}

    # ── Stage 3 + 4: EMBED + UPSERT in batches ───────────────────────────────
    # We send chunks to ChromaDB in batches of 64.
    # Inside each upsert() call:
    #   1. ChromaDB passes the batch of texts to embedding_fn (MiniLM)
    #   2. MiniLM encodes each text → 384-dim float32 vector
    #   3. ChromaDB stores (id, text, vector, metadata) in its HNSW index
    #
    # upsert vs add: upsert overwrites if the ID already exists — safe to re-run.
    batch_size = 64
    for i in range(0, len(all_docs), batch_size):
        batch_slice = slice(i, i + batch_size)
        collection.upsert(
            ids=all_ids[batch_slice],
            documents=all_docs[batch_slice],  # ChromaDB embeds these via MiniLM
            metadatas=all_metas[batch_slice],
        )
        print(f"    Upserted batch {i // batch_size + 1} "
              f"({min(i + batch_size, len(all_docs))}/{len(all_docs)} chunks)")

    final_count = collection.count()
    print(f"\n  ✓ {collection_name}: {final_count} total chunks stored  "
          f"(from {len(files)} files)")

    return {
        "collection": collection_name,
        "chunks":     final_count,
        "files":      len(files),
        "source_dir": str(source_dir),
    }


# ---------------------------------------------------------------------------
# Public query helper — used by agents at runtime
#
# At query time the pipeline is:
#   query_text (str)
#     → embedding_fn([query_text])          → 384-dim query vector
#     → ChromaDB HNSW ANN search            → top-N nearest chunk vectors
#     → return (text, metadata, distance)   → agent reads .text for context
# ---------------------------------------------------------------------------

def query_collection(
    collection_name: str,
    query_text: str,
    n_results: int = 5,
    where: Dict[str, Any] | None = None,
) -> List[Dict[str, Any]]:
    """
    Query a single named ChromaDB collection and return the top-N chunks.

    How retrieval works:
      1. The query_text is embedded with the same MiniLM model used at ingest.
      2. ChromaDB's HNSW index performs approximate nearest-neighbour (ANN)
         search over all stored vectors using cosine distance.
      3. The N chunks with the smallest cosine distance are returned.

    Args:
        collection_name: One of 'historical_precedents', 'export_control_corpus',
                         or 'india_sourcing_corpus'.
        query_text:      Natural-language question or keyword string.
        n_results:       How many chunks to return (capped at collection size).
        where:           Optional ChromaDB metadata filter, e.g.:
                           {"disruption_category": "geopolitical"}
                           {"severity_hint": {"$in": ["HIGH", "CRITICAL"]}}

    Returns:
        List of dicts with keys:
          "text"        — raw chunk text (pass this to the agent as context)
          "metadata"    — the metadata dict stored at ingest time
          "distance"    — cosine distance [0.0, 2.0]; lower = more relevant
          "collection"  — which collection this hit came from
    """
    from src.rag.utils import get_chroma_client, get_embedding_model

    client = get_chroma_client()
    embed_fn = get_embedding_model()

    try:
        col = client.get_collection(name=collection_name, embedding_function=embed_fn)
    except Exception:
        # Collection doesn't exist yet — return empty rather than crashing
        return []

    count = col.count()
    if count == 0:
        return []

    results = col.query(
        query_texts=[query_text],        # ChromaDB embeds this internally
        n_results=min(n_results, count), # can't request more than what's stored
        where=where,                     # optional metadata pre-filter
        include=["documents", "metadatas", "distances"],
    )

    # Unpack the nested list structure ChromaDB returns
    # (it supports batch queries, so each field is a list-of-lists)
    docs      = results.get("documents",  [[]])[0]
    metas     = results.get("metadatas",  [[]])[0]
    distances = results.get("distances",  [[]])[0]

    return [
        {
            "text":       doc,
            "metadata":   meta,
            "distance":   dist,
            "collection": collection_name,
        }
        for doc, meta, dist in zip(docs, metas, distances)
    ]


def query_multi_collection(
    query_text: str,
    collections: List[str] | None = None,
    n_per_collection: int = 3,
    where: Dict[str, Any] | None = None,
) -> List[Dict[str, Any]]:
    """
    Query multiple collections simultaneously and return merged results sorted
    by relevance (ascending cosine distance = most relevant first).

    Used by Agent 7 (Mitigation Recommendation), which needs context from
    all three collections: historical crises, export-control policy, and
    India-sourcing alternatives.

    Args:
        collections:       List of collection names to query. Defaults to all three.
        n_per_collection:  Results pulled from each collection before merging.

    Returns:
        Merged list of chunk dicts, sorted by distance ascending.
    """
    if collections is None:
        collections = list(COLLECTION_NAMES.values())

    all_results: List[Dict[str, Any]] = []
    for cname in collections:
        hits = query_collection(
            collection_name=cname,
            query_text=query_text,
            n_results=n_per_collection,
            where=where,
        )
        all_results.extend(hits)

    # Merge and re-rank across collections.
    # Lower cosine distance = more semantically similar to the query.
    all_results.sort(key=lambda x: x["distance"])
    return all_results


# ---------------------------------------------------------------------------
# CLI entry point
# ---------------------------------------------------------------------------

def main() -> None:
    """
    Command-line interface for building and smoke-testing the RAG collections.

    Examples:
      # Build / update all three collections (incremental — safe to re-run)
      python scripts/build_rag_collections.py

      # Wipe ChromaDB and rebuild everything from scratch
      python scripts/build_rag_collections.py --flush

      # Build only one collection
      python scripts/build_rag_collections.py --collection historical_precedents

      # Smoke-test query after building
      python scripts/build_rag_collections.py --query "Taiwan earthquake semiconductor"
    """
    parser = argparse.ArgumentParser(
        description="Build ChromaDB RAG collections for Capstone Project 8"
    )
    parser.add_argument(
        "--flush", action="store_true",
        help="Delete ALL existing ChromaDB data and rebuild from scratch. "
             "Use this when source files have changed significantly."
    )
    parser.add_argument(
        "--query", type=str, default=None,
        help="After building, run a smoke-test query against all collections "
             "to verify retrieval is working."
    )
    parser.add_argument(
        "--collection", type=str, default=None,
        choices=list(COLLECTION_NAMES.keys()),
        help="Build only a single named collection instead of all three."
    )
    args = parser.parse_args()

    # Load ChromaDB and the SentenceTransformer embedding function
    _import_chromadb()

    if args.flush and CHROMA_DIR.exists():
        print(f"[FLUSH] Removing existing ChromaDB at {CHROMA_DIR}")
        shutil.rmtree(CHROMA_DIR)
        reset_chroma_client()

    CHROMA_DIR.mkdir(parents=True, exist_ok=True)

    client = get_chroma_client()
    embed_fn = get_embedding_model()
    embed_model_label = resolve_embedding_model_name()

    # Print configuration so it's clear what we're doing
    print(f"\nEmbedding model : {embed_model_label}  (384-dim, cosine space)")
    print(f"ChromaDB path   : {CHROMA_DIR}")
    print(f"RAG data root   : {RAG_DATA_ROOT}")

    # Determine which collections to build
    if args.collection and args.collection in COLLECTION_NAMES:
        collections_to_build = {args.collection: COLLECTION_NAMES[args.collection]}
    else:
        collections_to_build = COLLECTION_NAMES

    summary: List[Dict[str, Any]] = []

    for cname in collections_to_build:
        source_dir = RAG_DATA_ROOT / cname
        chunk_cfg  = COLLECTION_CHUNK_CONFIG[cname]
        result = build_collection(
            collection_name=cname,
            source_dir=source_dir,
            chroma_client=client,
            embedding_fn=embed_fn,
            chunk_cfg=chunk_cfg,
        )
        summary.append(result)

    # Build summary table
    print("\n" + "=" * 60)
    print("  BUILD SUMMARY")
    print("=" * 60)
    total_chunks = 0
    total_files  = 0
    for r in summary:
        n_chunks = r["chunks"]
        n_files  = r.get("files", 0)
        print(f"  {r['collection']:35s}  {n_chunks:5d} chunks  ({n_files} files)")
        total_chunks += n_chunks
        total_files  += n_files
    print(f"  {'TOTAL':35s}  {total_chunks:5d} chunks  ({total_files} files)")
    print("=" * 60)
    print(f"\n  Embedding model:  {embed_model_label}")
    print(f"  Vector dims:      384  (float32)")
    print(f"  Distance metric:  cosine  (0.0 = identical, 2.0 = opposite)")
    print(f"  Storage:          {CHROMA_DIR}")

    # Optional smoke-test: embed a query and retrieve top hits from all collections
    if args.query:
        print(f"\n[SMOKE TEST] '{args.query}'\n")
        results = query_multi_collection(
            query_text=args.query,
            n_per_collection=2,
        )
        if not results:
            print("  No results — is the RAG data directory populated?")
        for hit in results[:6]:
            print(f"  [{hit['collection']}]  dist={hit['distance']:.3f}  "
                  f"file={hit['metadata'].get('source_file', '?')}")
            print(f"    {hit['text'][:200].strip()!r}")
            print()


if __name__ == "__main__":
    main()
