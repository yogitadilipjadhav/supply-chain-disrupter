from __future__ import annotations

import hashlib
import logging
import os
import re
import threading
from pathlib import Path
from typing import Any, Dict, List

import chromadb
import pandas as pd
from chromadb.utils.embedding_functions import SentenceTransformerEmbeddingFunction
from docx import Document as DocxDocument
from pypdf import PdfReader

from src.utils.etl_loader import EXCEL_SOURCE, read_excel_sheets

PROJECT_ROOT = Path(__file__).resolve().parents[2]
CHROMA_DIR = PROJECT_ROOT / "outputs" / "chromadb"
DEFAULT_COLLECTION_NAME = "electronics_supply_chain_knowledge"
PLAYBOOKS_DIR = PROJECT_ROOT / "config" / "playbooks"
STATIC_CONTEXT_DIR = PROJECT_ROOT / "data" / "raw" / "RAG_data"
EMBEDDING_MODEL = "all-MiniLM-L6-v2"
FINETUNED_EMBEDDING_MODEL = PROJECT_ROOT / "fine_tuning" / "models" / "supply_chain_embeddings"
CHUNK_SIZE = 1000
logger = logging.getLogger(__name__)
CHUNK_OVERLAP = 200

_EMBEDDING_WEIGHT_FILES = ("model.safetensors", "pytorch_model.bin")


def _embedding_weights_present(model_dir: Path) -> bool:
    """True when a local embedding directory contains loadable model weights."""
    return model_dir.is_dir() and any((model_dir / name).exists() for name in _EMBEDDING_WEIGHT_FILES)

# Module-level singleton — ChromaDB PersistentClient holds an exclusive file lock
# on chroma.sqlite3. Creating multiple instances in the same process causes WinError 32
# ("file is being used by another process"). A single shared client avoids this.
_chroma_client: chromadb.PersistentClient | None = None
_chroma_lock = threading.Lock()


def get_chroma_client() -> chromadb.PersistentClient:
    global _chroma_client
    with _chroma_lock:
        if _chroma_client is None:
            CHROMA_DIR.mkdir(parents=True, exist_ok=True)
            _chroma_client = chromadb.PersistentClient(path=str(CHROMA_DIR))
    return _chroma_client


def reset_chroma_client() -> None:
    """Drop the cached client (e.g. after --flush rebuild)."""
    global _chroma_client
    with _chroma_lock:
        _chroma_client = None


def resolve_embedding_model_name() -> str:
    """
    Resolve the sentence-transformers model id or local path for embeddings.

    Priority:
      1. EMBEDDING_MODEL_PATH env var (local directory or Hugging Face repo id)
      2. fine_tuning/models/supply_chain_embeddings/ (Phase A output)
      3. EMBEDDING_MODEL constant (base model)
    """
    custom = os.getenv("EMBEDDING_MODEL_PATH", "").strip()
    if custom:
        custom_path = Path(custom)
        if custom_path.exists():
            if _embedding_weights_present(custom_path):
                logger.info("Using custom embedding model (local): %s", custom)
                return str(custom_path)
            logger.warning(
                "EMBEDDING_MODEL_PATH %s exists but has no model weights — "
                "falling back to base model %s",
                custom,
                EMBEDDING_MODEL,
            )
        else:
            # Hugging Face repo ids (org/model) are not local paths — pass through.
            logger.info("Using custom embedding model: %s", custom)
            return custom
    if _embedding_weights_present(FINETUNED_EMBEDDING_MODEL):
        model_name = str(FINETUNED_EMBEDDING_MODEL)
        logger.info("Using fine-tuned embedding model (Phase A output): %s", model_name)
        return model_name
    if FINETUNED_EMBEDDING_MODEL.exists():
        logger.info(
            "Fine-tuned embedding dir present but no weights (Colab/git clone) — "
            "using base model: %s",
            EMBEDDING_MODEL,
        )
    logger.info("Using base embedding model (fine-tuned model not found): %s", EMBEDDING_MODEL)
    return EMBEDDING_MODEL


def get_embedding_model() -> SentenceTransformerEmbeddingFunction:
    """Return SentenceTransformerEmbeddingFunction for the resolved model."""
    return SentenceTransformerEmbeddingFunction(
        model_name=resolve_embedding_model_name(),
        normalize_embeddings=True,
    )


def _chunk_text(text: str) -> List[str]:
    text = re.sub(r"\r\n?", "\n", text).strip()
    if not text:
        return []
    if len(text) <= CHUNK_SIZE:
        return [text]

    chunks: List[str] = []
    start = 0
    while start < len(text):
        end = min(len(text), start + CHUNK_SIZE)
        if end < len(text):
            split = max(
                text.rfind("\n\n", start + CHUNK_SIZE // 2, end),
                text.rfind("\n", start + CHUNK_SIZE // 2, end),
                text.rfind(". ", start + CHUNK_SIZE // 2, end),
            )
            if split > start:
                end = split + 1
        chunk = text[start:end].strip()
        if chunk:
            chunks.append(chunk)
        if end >= len(text):
            break
        start = max(start + 1, end - CHUNK_OVERLAP)
    return chunks


def _doc_id(source: str, doc_type: str, key: str, chunk_index: int) -> str:
    value = f"{source}|{doc_type}|{key}|{chunk_index}".encode("utf-8")
    return hashlib.sha256(value).hexdigest()


def _add_document(
    documents: List[str],
    metadatas: List[Dict[str, Any]],
    ids: List[str],
    *,
    text: str,
    source: str,
    doc_type: str,
    key: str,
    metadata: Dict[str, Any] | None = None,
) -> None:
    for index, chunk in enumerate(_chunk_text(text)):
        documents.append(chunk)
        metadatas.append(
            {
                "source": source,
                "type": doc_type,
                "domain": "electronics_semiconductors",
                "dataset_owner": "Varun",
                "beauty_products_included": False,
                "chunk_index": index,
                **(metadata or {}),
            }
        )
        ids.append(_doc_id(source, doc_type, key, index))


def _workbook_documents(
    excel_path: Path,
) -> tuple[List[str], List[Dict[str, Any]], List[str], Dict[str, int]]:
    sheets = read_excel_sheets(excel_path)
    documents: List[str] = []
    metadatas: List[Dict[str, Any]] = []
    ids: List[str] = []
    counts = {
        "data_dictionary": 0,
        "workbook_context": 0,
        "event_profiles": 0,
        "mitigation_guidance": 0,
        "semiconductor_events": 0,
    }

    guide = sheets["Column Guide (Lite)"].dropna(how="all")
    for row in guide.to_dict(orient="records"):
        text = (
            f"Agent: {row.get('Agent')}\n"
            f"Field: {row.get('Column')}\n"
            f"Source type: {row.get('Source Type')}\n"
            f"Purpose: {row.get('Purpose')}"
        )
        _add_document(
            documents,
            metadatas,
            ids,
            text=text,
            source=excel_path.name,
            doc_type="data_dictionary",
            key=str(row.get("Column")),
            metadata={"agent": str(row.get("Agent")), "field": str(row.get("Column"))},
        )
        counts["data_dictionary"] += 1

    legend = sheets["Legend"].dropna(how="all")
    for index, row in legend.iterrows():
        text = f"{row.iloc[0]}: {row.iloc[1]}"
        _add_document(
            documents,
            metadatas,
            ids,
            text=text,
            source=excel_path.name,
            doc_type="workbook_context",
            key=str(index),
        )
        counts["workbook_context"] += 1

    master = sheets["Lite Master"]
    grouped = master.groupby("Disruption_Event_Label", dropna=False)
    for label, frame in grouped:
        label_text = "UNLABELLED" if pd.isna(label) else str(label)
        order_dates = pd.to_datetime(frame["Order_Date"], errors="coerce")
        top_regions = ", ".join(
            f"{name} ({count})"
            for name, count in frame["Order_Region"].value_counts().head(5).items()
        )
        top_products = ", ".join(
            f"{name} ({count})"
            for name, count in frame["Product_Name"].value_counts().head(5).items()
        )
        text = (
            f"Electronics disruption profile: {label_text}\n"
            f"Order observations: {len(frame):,}\n"
            f"Date range: {order_dates.min():%Y-%m-%d} to "
            f"{order_dates.max():%Y-%m-%d}\n"
            f"Top affected regions: {top_regions}\n"
            f"Top products: {top_products}\n"
            f"Average composite risk: {frame['Risk_Score_Composite'].mean():.3f}\n"
            f"Average supply disruption index: "
            f"{frame['Supply_Disruption_Index'].mean():.3f}\n"
            f"Average lead-time variance: "
            f"{frame['Lead_Time_Variance_Days'].mean():.2f} days\n"
            f"Average stockout probability: "
            f"{frame['Stockout_Probability_Pct'].mean():.2f}%\n"
            f"Alternate supplier availability: "
            f"{frame['Alternate_Supplier_Available'].mean() * 100:.1f}%"
        )
        _add_document(
            documents,
            metadatas,
            ids,
            text=text,
            source=excel_path.name,
            doc_type="event_profile",
            key=label_text,
            metadata={"event_label": label_text, "observations": len(frame)},
        )
        counts["event_profiles"] += 1

    recommendations = (
        master[
            ["Disruption_Event_Label", "Mitigation_Recommendation"]
        ]
        .dropna()
        .drop_duplicates()
    )
    for _, row in recommendations.iterrows():
        label = str(row["Disruption_Event_Label"])
        recommendation = str(row["Mitigation_Recommendation"])
        text = (
            f"Electronics mitigation guidance\n"
            f"Risk label: {label}\n"
            f"Recommended response: {recommendation}"
        )
        _add_document(
            documents,
            metadatas,
            ids,
            text=text,
            source=excel_path.name,
            doc_type="mitigation_guidance",
            key=f"{label}|{recommendation}",
            metadata={"event_label": label},
        )
        counts["mitigation_guidance"] += 1

    signals = sheets["Semiconductor Signals"].dropna(how="all")
    event_rows = signals[
        signals["Known Disruption Event"].notna()
        & ~signals["Known Disruption Event"].astype(str).isin(["—", "-", "None"])
    ]
    event_rows = event_rows.drop_duplicates(
        subset=["Year", "Country", "Company", "Known Disruption Event"]
    )
    for _, row in event_rows.iterrows():
        event = str(row["Known Disruption Event"])
        text = (
            f"Historical semiconductor disruption signal\n"
            f"Year: {row['Year']}\nCountry: {row['Country']}\n"
            f"Company: {row['Company']}\nEvent: {event}\n"
            f"Known severity: {row['Known Severity']}\n"
            f"Supply disruption index: {row['Supply Disruption Index']}\n"
            f"Semiconductor security risk: "
            f"{row['Semiconductor Security Risk']}\n"
            f"Natural disaster risk: {row['Natural Disaster Risk']}\n"
            f"Factory shutdown risk: {row['Factory Shutdown Risk']}\n"
            f"Export control level: {row['Export Control Level']}\n"
            f"Chip price index: {row['Chip Price Index']}"
        )
        key = f"{row['Year']}|{row['Country']}|{row['Company']}|{event}"
        _add_document(
            documents,
            metadatas,
            ids,
            text=text,
            source=excel_path.name,
            doc_type="semiconductor_event",
            key=key,
            metadata={
                "year": int(row["Year"]),
                "country": str(row["Country"]),
                "company": str(row["Company"]),
                "event": event,
                "severity": str(row["Known Severity"]),
            },
        )
        counts["semiconductor_events"] += 1

    return documents, metadatas, ids, counts


def _playbook_documents(
    playbooks_dir: Path,
) -> tuple[List[str], List[Dict[str, Any]], List[str], int]:
    documents: List[str] = []
    metadatas: List[Dict[str, Any]] = []
    ids: List[str] = []
    count = 0
    if not playbooks_dir.exists():
        return documents, metadatas, ids, count

    for path in sorted(playbooks_dir.glob("*.txt")):
        text = path.read_text(encoding="utf-8")
        _add_document(
            documents,
            metadatas,
            ids,
            text=text,
            source=path.name,
            doc_type="mitigation_playbook",
            key=path.stem,
            metadata={"title": path.stem.replace("_", " ")},
        )
        count += 1
    return documents, metadatas, ids, count


def _docx_text(path: Path) -> str:
    document = DocxDocument(path)
    sections: List[str] = []

    paragraphs = [
        paragraph.text.strip()
        for paragraph in document.paragraphs
        if paragraph.text.strip()
    ]
    if paragraphs:
        sections.append("\n\n".join(paragraphs))

    for table_index, table in enumerate(document.tables, start=1):
        rows = []
        for row in table.rows:
            values = [cell.text.strip() for cell in row.cells]
            if any(values):
                rows.append(" | ".join(values))
        if rows:
            sections.append(f"Table {table_index}\n" + "\n".join(rows))

    return "\n\n".join(sections)


def _static_context_documents(
    context_dir: Path,
) -> tuple[List[str], List[Dict[str, Any]], List[str], Dict[str, int]]:
    documents: List[str] = []
    metadatas: List[Dict[str, Any]] = []
    ids: List[str] = []
    counts = {
        "static_context_files": 0,
        "static_pdf_files": 0,
        "static_pdf_pages": 0,
        "static_docx_files": 0,
    }
    if not context_dir.exists():
        return documents, metadatas, ids, counts

    for path in sorted(context_dir.glob("*.pdf")):
        reader = PdfReader(path)
        indexed_pages = 0
        for page_number, page in enumerate(reader.pages, start=1):
            text = page.extract_text() or ""
            if not text.strip():
                continue
            _add_document(
                documents,
                metadatas,
                ids,
                text=text,
                source=path.name,
                doc_type="static_report",
                key=f"{path.stem}|page-{page_number}",
                metadata={
                    "title": path.stem.replace("_", " "),
                    "file_type": "pdf",
                    "page": page_number,
                },
            )
            indexed_pages += 1

        if indexed_pages:
            counts["static_context_files"] += 1
            counts["static_pdf_files"] += 1
            counts["static_pdf_pages"] += indexed_pages

    for path in sorted(context_dir.glob("*.docx")):
        text = _docx_text(path)
        if not text.strip():
            continue
        _add_document(
            documents,
            metadatas,
            ids,
            text=text,
            source=path.name,
            doc_type="static_report",
            key=path.stem,
            metadata={
                "title": path.stem.replace("_", " "),
                "file_type": "docx",
            },
        )
        counts["static_context_files"] += 1
        counts["static_docx_files"] += 1

    return documents, metadatas, ids, counts


def build_chroma_from_default_excel(flush_existing: bool = False) -> Dict[str, Any]:
    return build_chroma_complete(flush_existing=flush_existing)


def build_rag_corpus_complete(flush_existing: bool = True) -> Dict[str, Any]:
    return build_chroma_complete(flush_existing=flush_existing)


def build_chroma_complete(
    flush_existing: bool = True,
    excel_path: Path = EXCEL_SOURCE,
    playbooks_dir: Path = PLAYBOOKS_DIR,
    static_context_dir: Path = STATIC_CONTEXT_DIR,
) -> Dict[str, Any]:
    """Build one electronics semantic store from all committed static sources."""
    if flush_existing:
        # Drop the collection via the API rather than deleting the directory.
        # shutil.rmtree on a live PersistentClient causes WinError 32 on Windows
        # because the client holds an exclusive lock on chroma.sqlite3.
        try:
            get_chroma_client().delete_collection(name=DEFAULT_COLLECTION_NAME)
        except Exception:
            pass  # collection may not exist yet — that's fine

    workbook_docs, workbook_meta, workbook_ids, counts = _workbook_documents(
        excel_path
    )
    playbook_docs, playbook_meta, playbook_ids, playbook_count = (
        _playbook_documents(playbooks_dir)
    )
    context_docs, context_meta, context_ids, context_counts = (
        _static_context_documents(static_context_dir)
    )

    documents = workbook_docs + playbook_docs + context_docs
    metadatas = workbook_meta + playbook_meta + context_meta
    ids = workbook_ids + playbook_ids + context_ids
    if not documents:
        raise ValueError("No electronics knowledge documents were generated")

    client = get_chroma_client()
    collection = client.get_or_create_collection(
        name=DEFAULT_COLLECTION_NAME,
        embedding_function=get_embedding_model(),
        metadata={
            "hnsw:space": "cosine",
            "domain": "electronics_semiconductors",
            "dataset_owner": "Varun",
        },
    )

    batch_size = 100
    for start in range(0, len(documents), batch_size):
        collection.upsert(
            ids=ids[start : start + batch_size],
            documents=documents[start : start + batch_size],
            metadatas=metadatas[start : start + batch_size],
        )

    return {
        "collection": DEFAULT_COLLECTION_NAME,
        "chunks": collection.count(),
        "source_documents": (
            sum(counts.values())
            + playbook_count
            + context_counts["static_context_files"]
        ),
        "playbooks": playbook_count,
        **counts,
        **context_counts,
        "domain": "electronics_semiconductors",
        "dataset_owner": "Varun",
        "beauty_products_included": False,
        "embedding_model": resolve_embedding_model_name(),
    }


def query_chroma_rag(
    query: str,
    n_results: int = 5,
    collection_name: str = DEFAULT_COLLECTION_NAME,
    where: Dict[str, Any] | None = None,
) -> List[Dict[str, Any]]:
    client = get_chroma_client()
    try:
        collection = client.get_collection(
            name=collection_name,
            embedding_function=get_embedding_model(),
        )
    except Exception:
        return []

    count = collection.count()
    if count == 0:
        return []
    results = collection.query(
        query_texts=[query],
        n_results=min(n_results, count),
        where=where,
        include=["documents", "metadatas", "distances"],
    )

    documents = results.get("documents", [[]])[0]
    metadatas = results.get("metadatas", [[]])[0]
    distances = results.get("distances", [[]])[0]
    return [
        {"text": doc, "metadata": metadata, "distance": distance}
        for doc, metadata, distance in zip(documents, metadatas, distances)
    ]
